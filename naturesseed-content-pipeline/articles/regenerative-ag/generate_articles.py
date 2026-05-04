"""
Generate .docx files for Nature's Seed regenerative ag article series.
Run: python3 generate_articles.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_font(run, name, size, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level):
    """Add a styled heading using Noto Serif Display style fallback."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    sizes = {1: 22, 2: 16, 3: 14}
    set_font(run, "Noto Serif Display", sizes.get(level, 14), bold=True, color=(29, 67, 50))
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_body(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, "Inter", 11, color=(33, 37, 41))
    para.paragraph_format.space_after = Pt(8)
    return para


def add_meta_block(doc, category, read_time, meta_desc):
    """Add category / read time / meta description block."""
    para = doc.add_paragraph()
    run = para.add_run(f"Category: {category}  |  {read_time}  |  For internal use")
    set_font(run, "Inter", 9, color=(73, 80, 87))
    para.paragraph_format.space_after = Pt(4)

    para2 = doc.add_paragraph()
    run2 = para2.add_run(f"META: {meta_desc}")
    set_font(run2, "Inter", 9, bold=True, color=(44, 106, 79))
    para2.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()  # spacer


def add_divider(doc):
    para = doc.add_paragraph()
    run = para.add_run("─" * 60)
    set_font(run, "Inter", 9, color=(200, 200, 200))
    para.paragraph_format.space_after = Pt(4)


def _set_cell_shading(cell, hex_color):
    """Set background fill color on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows):
    """Add a styled table. headers = list of str, rows = list of list of str."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        para = hdr_cells[i].paragraphs[0]
        para.clear()
        run = para.add_run(h)
        run.font.name = "Inter"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(hdr_cells[i], "2D6A4F")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        fill = "F8F9FA" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, cell_text in enumerate(row_data[:len(headers)]):
            para = row_cells[c_idx].paragraphs[0]
            para.clear()
            run = para.add_run(cell_text)
            run.font.name = "Inter"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(33, 37, 41)
            _set_cell_shading(row_cells[c_idx], fill)

    doc.add_paragraph()  # spacing after table


def add_references_section(doc, refs):
    """Add an APA references section. refs = list of str (formatted citations)."""
    if not refs:
        return
    add_divider(doc)
    add_heading(doc, "References", 2)
    for i, ref in enumerate(refs, 1):
        para = doc.add_paragraph()
        num_run = para.add_run(f"[{i}] ")
        set_font(num_run, "Inter", 9, bold=True, color=(44, 106, 79))
        ref_run = para.add_run(ref)
        set_font(ref_run, "Inter", 9, color=(73, 80, 87))
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.left_indent = Inches(0.25)


def build_doc(article):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(article["title"])
    set_font(title_run, "Noto Serif Display", 26, bold=True, color=(27, 67, 50))
    title_para.paragraph_format.space_after = Pt(8)

    add_meta_block(doc, article["category"], article["read_time"], article["meta"])
    add_divider(doc)
    doc.add_paragraph()

    # Body parsing
    for block in article["body"]:
        kind = block["type"]
        text = block["text"]

        if kind == "h2":
            add_heading(doc, text, 2)
        elif kind == "h3":
            add_heading(doc, text, 3)
        elif kind == "p":
            add_body(doc, text)
        elif kind == "bullet":
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(text)
            set_font(run, "Inter", 11, color=(33, 37, 41))
            para.paragraph_format.space_after = Pt(4)
        elif kind == "cta":
            para = doc.add_paragraph()
            run = para.add_run(text)
            set_font(run, "Inter", 11, bold=True, color=(201, 106, 46))
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(8)
        elif kind == "table":
            add_table(doc, block["headers"], block["rows"])
        elif kind == "note":
            para = doc.add_paragraph()
            run = para.add_run(text)
            set_font(run, "Inter", 10, color=(73, 80, 87))
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.left_indent = Inches(0.25)

    if article.get("references"):
        add_references_section(doc, article["references"])

    return doc


# ---------------------------------------------------------------------------
# ARTICLE DATA
# ---------------------------------------------------------------------------

ARTICLES = [

# ── ARTICLE 1 ──────────────────────────────────────────────────────────────
{
  "title": "What Is Regenerative Agriculture? A Plain-English Guide",
  "category": "Foundations",
  "read_time": "6 min read",
  "meta": "Regenerative agriculture for small farms explained simply — what it actually means, what it costs, and how to start with one practice on your land.",
  "slug": "01-what-is-regenerative-agriculture",
  "body": [
    {"type":"p","text":"You've probably heard \"regenerative agriculture\" tossed around at farm shows, in seed catalogs, and all over social media. It sounds expensive. It sounds complicated. And it sounds like something designed for large operations with a full-time agronomist on staff."},
    {"type":"p","text":"It's not."},
    {"type":"p","text":"Regenerative agriculture for small farms is simpler than the marketing makes it seem. At its core, it's a set of farming practices aimed at leaving your land in better shape than you found it — better soil, more water retention, lower input costs over time. That's it."},
    {"type":"p","text":"Let's cut through the noise."},
    {"type":"h2","text":"It's Not Organic (But It Overlaps)"},
    {"type":"p","text":"A lot of people assume regenerative and organic are the same thing. They're not."},
    {"type":"p","text":"Organic farming is defined by what you don't use — no synthetic pesticides, no synthetic fertilizers. It's a certification with specific rules. You can farm organically and still plow your fields every year, leave soil bare between seasons, and run a monoculture. Technically organic. Not regenerative."},
    {"type":"p","text":"Regenerative agriculture is defined by what you're building. The focus is on soil health, biological activity, and long-term resilience. Some regenerative farmers are certified organic. Many aren't. You can use regenerative practices without ever pursuing organic certification."},
    {"type":"p","text":"The overlap? Both approaches tend to reduce chemical inputs over time. But regenerative ag gets there by building a system that needs less, not by following a checklist."},
    {"type":"h2","text":"The Three Things Regenerative Ag Is Actually Trying to Do"},
    {"type":"p","text":"Strip away the philosophy and you're left with three practical goals."},
    {"type":"h3","text":"1. Build Soil Organic Matter"},
    {"type":"p","text":"Organic matter is the dark, spongy stuff in healthy soil. It holds water, feeds soil biology, and makes nutrients available to plants. Most farmed soils are depleted — often below 2% organic matter. Regenerative practices push that number up over time. Even a 1% increase in organic matter helps your soil hold significantly more water per acre."},
    {"type":"h3","text":"2. Close Nutrient Cycles"},
    {"type":"p","text":"Conventional farming imports nutrients (fertilizer), grows a crop, exports the crop, and repeats. Nutrients leave the farm and don't come back. Regenerative systems try to cycle nutrients on-farm — through cover crops that fix nitrogen, through livestock manure, through plant residue decomposing in place. You're not shipping your fertility off to someone else's land."},
    {"type":"h3","text":"3. Reduce Inputs Over Time"},
    {"type":"p","text":"This is the one that should get every budget-minded farmer's attention. A healthy, biologically active soil requires less fertilizer, less herbicide, less irrigation. You're not there on day one. But that's the direction you're heading. Every year the soil improves, your input costs have room to drop."},
    {"type":"h2","text":"What It Looks Like on a Small Operation"},
    {"type":"p","text":"On a 50-acre farm or ranch, regenerative agriculture isn't some grand redesign. It's a handful of specific practices layered in over time."},
    {"type":"p","text":"Cover crops. Planting something between your cash crops or after hay harvest keeps living roots in the ground and prevents bare soil. Legumes like clover and hairy vetch fix atmospheric nitrogen. Grasses add organic matter. Brassicas break up compaction. A simple two- or three-species mix after your main crop can start improving your soil biology immediately."},
    {"type":"p","text":"Rotational grazing. If you run livestock, moving them frequently — rather than leaving them on the same pasture all season — gives grass time to recover. Grass that recovers builds deeper roots. Deeper roots deposit more carbon in the soil. Well-managed grazing is one of the fastest ways to build soil organic matter."},
    {"type":"p","text":"Reducing tillage. Every time you plow or disk, you disrupt the fungal networks and microbial communities that make healthy soil function. Reducing tillage passes lets the biology rebuild. This doesn't mean no-till is required from day one. It means asking whether every tillage pass is earning its keep."},
    {"type":"h2","text":"The Budget Reality"},
    {"type":"p","text":"Here's what nobody tells you upfront: you don't have to do all of this at once."},
    {"type":"p","text":"Buying a new no-till drill, hiring a grazing consultant, and replanting every acre in diverse cover crop mixes in year one would cost serious money. That's not the point."},
    {"type":"p","text":"The point is to start one practice on one field and learn from it. Cover crop seed for a 20-acre field costs a few hundred dollars. Splitting one large pasture with a single strand of temporary electric fence to start a basic rotation is cheap. Neither of these requires a loan."},
    {"type":"p","text":"Most farmers who've gone down this road say the same thing: they started with one thing, saw a result, and added the next practice when they were ready. Patience is the main input."},
    {"type":"h2","text":"Where to Start Without Hiring a Consultant"},
    {"type":"p","text":"You don't need to hire anyone. Here's a simple sequence:"},
    {"type":"bullet","text":"Get a soil test first. A standard soil test from your county extension office costs $15–30. It tells you your pH, organic matter, and major nutrients. This is your baseline."},
    {"type":"bullet","text":"Pick one regenerative practice. If you grow row crops, plant a cover crop after harvest. If you run cattle, try a simple two-paddock rotation."},
    {"type":"bullet","text":"Track it. Take photos. Note which fields you changed. Run another soil test in two to three years and compare."},
    {"type":"bullet","text":"Use what's available. NRCS offers cost-share programs for cover cropping and rotational grazing fencing. Your local office can tell you what's available. It's free money most small landowners leave on the table."},
    {"type":"cta","text":"→ Nature's Seed offers cover crop and pasture seed mixes designed for real-world farm conditions. A practical starting point whether you're seeding a small test plot or covering a few hundred acres. Browse cover crop seed at naturesseed.com."},
  ]
},

# ── ARTICLE 2 ──────────────────────────────────────────────────────────────
{
  "title": "The 5 Principles of Soil Health: Gabe Brown's Framework for Skeptical Ranchers",
  "category": "Foundations",
  "read_time": "7 min read",
  "meta": "Gabe Brown's 5 soil health principles for ranchers — practical, no-fluff breakdown of how a nearly-broke North Dakota rancher rebuilt his land without inputs.",
  "slug": "02-5-principles-soil-health",
  "body": [
    {"type":"p","text":"Gabe Brown doesn't look like a farming celebrity. He's a rancher from Burleigh County, North Dakota, who got hit with four consecutive crop failures in the late 1990s. Hail. Drought. Ice storms. He couldn't afford inputs. He had to figure something else out."},
    {"type":"p","text":"What he figured out became one of the most practical frameworks in modern agriculture: five principles of soil health that any rancher or small farmer can apply, regardless of operation size or budget."},
    {"type":"p","text":"These aren't theory. They came from a man who had no choice but to make them work. The soil health principles for ranchers that Brown developed have since been adopted by the NRCS, studied by soil scientists, and replicated on farms across the country."},
    {"type":"h2","text":"Principle 1 — Limit Disturbance"},
    {"type":"p","text":"Disturbance comes in three forms: mechanical, chemical, and biological."},
    {"type":"p","text":"Mechanical disturbance is tillage. Every time you run a plow or a disk through a field, you're destroying the fungal networks — mycorrhizal fungi — that connect plants to nutrients and water. You're also oxidizing organic matter, which releases carbon into the air instead of keeping it in the soil. Tillage has a place in some situations, but every pass should earn its keep."},
    {"type":"p","text":"Chemical disturbance is more nuanced. Herbicides, fungicides, and high-rate synthetic fertilizers can suppress or alter soil biology. This doesn't mean you can never use them. It means understanding that each application has a biological cost, and over-reliance keeps you on a treadmill."},
    {"type":"p","text":"Biological disturbance is overgrazing. Leaving livestock on the same ground too long removes the plant canopy, compacts soil, and gives invasive species a foothold. Brown considers overgrazing as damaging as tillage — maybe more so on rangelands."},
    {"type":"p","text":"Start here: identify your highest-disturbance practice and ask if it's necessary at current rates."},
    {"type":"h2","text":"Principle 2 — Keep Soil Covered"},
    {"type":"p","text":"Bare soil is a liability. It loses moisture. It erodes in wind and rain. It heats up to temperatures that kill surface biology. It's an open invitation for weeds."},
    {"type":"p","text":"Brown's rule is simple: something should be covering the ground at all times. That means living plants, crop residue, or mulch. Cover crops are the primary tool here. Planted after cash crop harvest or between grazing rotations, they provide immediate soil cover while also contributing to the next principle."},
    {"type":"p","text":"Even a simple winter rye or winter wheat seeding after harvest keeps your soil protected through the cold months and builds organic matter as it decomposes in spring. The economic argument is straightforward: every inch of topsoil you keep is money you didn't have to replace with fertilizer."},
    {"type":"h2","text":"Principle 3 — Maintain Living Roots Year-Round"},
    {"type":"p","text":"Plants feed soil biology through their roots. They release sugars and other compounds — called root exudates — that feed bacteria and fungi. Those organisms, in turn, make nutrients available to the plant. It's a two-way exchange."},
    {"type":"p","text":"When the ground is bare, that exchange stops. The biology goes dormant or dies back. You have to restart it every season, which takes time and inputs."},
    {"type":"p","text":"Living roots year-round means designing your system so something is always growing. For crop farmers, that means cover crops filling every gap in the calendar. For ranchers, it means managing grazing so pastures have adequate rest and root systems stay intact."},
    {"type":"p","text":"Perennial species are your friend here. Well-managed perennial pastures — diverse mixes of grasses and legumes — maintain living roots continuously without replanting. Clovers are cheap to establish and provide biological nitrogen fixation alongside continuous root activity. A pasture that includes legumes is feeding the soil while it feeds your livestock."},
    {"type":"h2","text":"Principle 4 — Maximize Diversity"},
    {"type":"p","text":"Monocultures underperform biologically diverse systems. Different plant species have different root architectures. Some go deep and break compaction. Some stay shallow and add organic matter near the surface. Different species host different microbial communities. A diverse plant community above ground means a diverse, resilient biology below ground."},
    {"type":"p","text":"Diversity applies to cover crops, pasture mixes, and cropping systems alike. A cover crop mix with five to seven species outperforms a single-species planting on almost every biological metric — nitrogen fixation, organic matter addition, weed suppression, and forage quality."},
    {"type":"p","text":"You don't have to do all of this at once. Start by adding one or two species to whatever cover crop or pasture seeding you're already doing."},
    {"type":"h2","text":"Principle 5 — Integrate Livestock"},
    {"type":"p","text":"This is the principle most crop farmers skip, and it's often the missing link. Livestock, managed well, accelerate every other principle. They trample residue into the soil surface, speeding decomposition. Their manure introduces diverse biology. Their hooves create small depressions that capture water and seed. And their grazing, when timed right, stimulates plant growth rather than suppressing it."},
    {"type":"p","text":"Brown argues that the separation of crop and livestock production — a 20th century efficiency move — is one of the root causes of soil degradation. Bringing animals back onto crop ground, even briefly, changes the biological trajectory of that soil."},
    {"type":"p","text":"For ranchers who already have livestock, this principle is about how you graze, not whether you graze. High-density, short-duration grazing with long rest periods mimics the movement of wild herds and produces dramatically different results than continuous grazing on the same ground."},
    {"type":"h2","text":"How to Apply This on a Budget"},
    {"type":"p","text":"You don't implement all five principles in year one. Here's a reasonable sequence for a small operation with limited capital:"},
    {"type":"bullet","text":"Year 1: Start with Principle 2 (soil cover). Plant a simple, affordable cover crop on your most vulnerable ground. Winter rye is cheap, reliable, and works almost everywhere. This costs seed money and nothing else."},
    {"type":"bullet","text":"Year 2: Add Principle 3 (living roots). Extend your cover crop calendar. Look at where you have bare ground in late summer or early fall and fill it."},
    {"type":"bullet","text":"Year 3: Work on Principle 4 (diversity). Upgrade from a single-species cover to a mix. Add a legume to your pasture seeding."},
    {"type":"bullet","text":"Year 4 and beyond: Tackle Principle 1 (reduce disturbance) and Principle 5 (integrate livestock) as your system matures."},
    {"type":"cta","text":"→ Nature's Seed offers cover crop seed mixes and pasture seed designed for real operational conditions — diverse, regionally appropriate options that make it easier to hit multiple principles at once. Browse at naturesseed.com."},
  ]
},

# ── ARTICLE 3 ──────────────────────────────────────────────────────────────
{
  "title": "What Your Soil Test Actually Tells You (And What Most Farmers Miss)",
  "category": "Foundations",
  "read_time": "5 min read",
  "meta": "Learn how to read a soil test — which numbers actually drive yield, what pH really controls, and why organic matter is the most overlooked result on the page.",
  "slug": "03-what-your-soil-test-tells-you",
  "body": [
    {"type":"p","text":"Most farmers run a soil test every few years. They get a report back, look at the phosphorus and potassium numbers, order some fertilizer, and call it done."},
    {"type":"p","text":"That's using about 20% of the information on the page."},
    {"type":"p","text":"Knowing how to read a soil test — really read it — gives you a picture of what's limiting your land and where your money will actually move the needle."},
    {"type":"h2","text":"The Numbers That Actually Drive Yield"},
    {"type":"p","text":"A standard soil test typically reports: pH, organic matter (OM), cation exchange capacity (CEC), phosphorus (P), potassium (K), and sometimes secondary and micronutrients like sulfur, calcium, and magnesium."},
    {"type":"p","text":"Three of those numbers are foundational. Everything else builds on them: pH controls nutrient availability across the board. Organic matter drives biological activity and water-holding capacity. CEC tells you how well your soil holds onto the nutrients you're applying."},
    {"type":"p","text":"Get those three right, and your fertilizer dollars go further. Ignore them, and you can throw inputs at the ground all season and wonder why yields aren't moving."},
    {"type":"h2","text":"What pH Really Controls"},
    {"type":"p","text":"pH is reported on a scale of 1 to 14, with 7 being neutral. Most crops want a pH between 6.0 and 7.0. Grass pastures are a little more forgiving. Legumes like clover need pH above 6.2 to fix nitrogen effectively."},
    {"type":"p","text":"Here's the part most farmers miss: pH doesn't just tell you whether your soil is acidic or alkaline. It controls the availability of nutrients that are already in your soil."},
    {"type":"p","text":"Phosphorus becomes chemically locked up at pH below 6.0 and above 7.5. You could have plenty of phosphorus in your soil and still have plants that can't access it. You're not low on phosphorus. You're low on pH management."},
    {"type":"p","text":"The same is true for most micronutrients. Iron, manganese, zinc, and boron are all more available in the 6.0–7.0 range. Lime is the common fix for low pH. It's one of the best returns on investment in agriculture — cheap, long-lasting, and it unlocks fertilizer you've already paid for. If your pH is below 6.0, fix that before you spend money on anything else."},
    {"type":"h2","text":"The Organic Matter Number Nobody Talks About"},
    {"type":"p","text":"Organic matter (OM) is reported as a percentage. A degraded agricultural soil might be at 1–2%. A healthy, biologically active soil is often 3–5% or higher."},
    {"type":"p","text":"Here's why it matters: every 1% increase in organic matter allows your soil to hold an additional 20,000 gallons of water per acre. In a dry year, that's the difference between a stressed crop and a productive one."},
    {"type":"p","text":"The other reason to pay attention to your OM number is that it's a direct measure of whether your management is moving in the right direction. If you're adding cover crops, reducing tillage, and managing grazing well, your organic matter should trend upward over time."},
    {"type":"p","text":"A good target for most Midwest and Great Plains soils is 3% or above. If you're below 2%, building organic matter should be your primary goal — ahead of fertilizer optimization."},
    {"type":"h2","text":"Phosphorus and Potassium — Stop Chasing Big Numbers"},
    {"type":"p","text":"Most university extension recommendations are calibrated for maximum yield under conventional management. They're designed to ensure you never limit yield due to nutrient deficiency. But they're not optimized for cost efficiency, and they don't account for biological nutrient cycling."},
    {"type":"p","text":"In a healthy, biologically active soil, phosphorus availability is enhanced by mycorrhizal fungi. Potassium is cycled through plant residue and microbial activity. You often don't need the levels a conventional recommendation calls for if your biology is functioning well."},
    {"type":"p","text":"Practically speaking: if your P and K test at medium or above, hold your fertilizer application and focus your resources on pH and organic matter. You'll likely see better results for less money."},
    {"type":"h2","text":"What the Test Can't Tell You"},
    {"type":"p","text":"A standard soil test is a chemical snapshot. It tells you what's there. It doesn't tell you:"},
    {"type":"bullet","text":"Biological activity — microbial and fungal populations that drive nutrient cycling don't show up on a standard test. A Haney test or PLFA analysis can give you a window into this."},
    {"type":"bullet","text":"Compaction — a penetrometer test or a screwdriver-in-the-ground test gives you more useful data than any lab report."},
    {"type":"bullet","text":"Water infiltration — pour a gallon of water on bare soil and time how long it takes to absorb. That observation tells you something no test captures."},
    {"type":"h2","text":"Your Next Move After the Test"},
    {"type":"bullet","text":"Fix pH first if you're outside the 6.0–7.0 range. Lime pays for itself."},
    {"type":"bullet","text":"Plan for organic matter if you're below 3%. Cover crops are your most cost-effective tool."},
    {"type":"bullet","text":"Address genuine nutrient deficiencies once pH is corrected. You may need less than the report suggests."},
    {"type":"bullet","text":"Repeat the test in two to three years and track your trend."},
    {"type":"cta","text":"→ Cover crops are one of the most powerful tools for building organic matter year over year. Nature's Seed carries cover crop seed options suited for a range of goals — from straight legume plantings to diverse multi-species mixes. Browse at naturesseed.com."},
  ]
},

# ── ARTICLE 4 ──────────────────────────────────────────────────────────────
{
  "title": "Cover Crops 101: How to Pick Your First Mix Without an Agronomy Degree",
  "category": "Foundations",
  "read_time": "8 min read",
  "meta": "Choosing the right cover crop mix for small farms doesn't have to be complicated. Practical guide to species, seeding rates, and termination for beginners.",
  "slug": "04-cover-crops-101-first-mix",
  "body": [
    {"type":"p","text":"Cover crop selection has gotten overcomplicated. Walk into any agricultural forum and you'll find people arguing about seeding rates to the ounce, specific species combinations for their exact soil type, and termination windows timed to the hour. It's enough to make a first-timer bail before they even start."},
    {"type":"p","text":"Here's the reality: a simple cover crop mix planted at a reasonable rate and terminated before it sets seed will improve your soil. You don't need to nail every detail. You need to start."},
    {"type":"h2","text":"Why a Mix Beats a Single Species"},
    {"type":"p","text":"When cover crops first got popular, a lot of farmers planted straight winter rye. It's cheap, it grows, it covers the ground. Not a bad start. But a diverse cover crop mix for small farms outperforms a single species on almost every measure."},
    {"type":"bullet","text":"Functional diversity means more soil benefits. Grasses add biomass and organic matter. Legumes fix nitrogen. Brassicas break up compaction. A mix does all three at once."},
    {"type":"bullet","text":"Diverse plantings support more soil biology. Different root architectures host different microbial communities. A field with five species supports a richer soil food web than a monoculture cover crop."},
    {"type":"bullet","text":"Risk management. If one species fails, the others carry the mix. Single-species plantings are all-or-nothing."},
    {"type":"p","text":"You don't need twelve species in a mix to see these benefits. Three to five well-chosen species is plenty for a first planting."},
    {"type":"h2","text":"The Three Functional Groups You Need"},
    {"type":"h3","text":"Grasses"},
    {"type":"p","text":"Grasses are your biomass producers. They grow fast, produce a lot of above- and below-ground material, and add significant organic matter as they break down. Common options: Winter rye (cheapest and most reliable cool-season grass, works almost anywhere), oats (good companion species, winter-kills in most climates which simplifies termination), sorghum-sudan (excellent warm-season option for summer cover, grows fast and produces massive biomass)."},
    {"type":"h3","text":"Legumes"},
    {"type":"p","text":"Legumes fix atmospheric nitrogen through root-zone bacteria. That nitrogen becomes available to your next crop as the cover breaks down. Common options: Crimson clover (one of the easiest legumes to establish, 60–120 lbs of nitrogen per acre under good conditions), hairy vetch (winter-hardy with excellent nitrogen fixation, pairs well with winter rye), field peas (fast-growing, palatable for grazing, good nitrogen contributor)."},
    {"type":"h3","text":"Brassicas"},
    {"type":"p","text":"Brassicas produce deep taproots that physically break up compaction layers. They also scavenge nutrients from lower soil horizons and bring them to the surface as they decompose. Common options: daikon radish (deep roots die in winter in most climates, leaving channels that improve drainage), turnip (similar function to radish, often more winter-hardy), rapeseed/canola (good biomass producer, works well in mixes)."},
    {"type":"h2","text":"How to Build Your Mix by Goal"},
    {"type":"p","text":"Don't just plant random species. Pick a goal, then build around it."},
    {"type":"bullet","text":"Suppress weeds: Lead with high-biomass species — winter rye, sorghum-sudan, or oats. Add a legume for nitrogen benefit. Suggested mix: Winter rye + hairy vetch."},
    {"type":"bullet","text":"Build organic matter: Maximize biomass. Grasses and brassicas are your primary drivers. Let them get as large as possible before terminating. Suggested mix: Oats + winter rye + turnip."},
    {"type":"bullet","text":"Fix nitrogen: Lead with legumes. Suggested mix: Hairy vetch + crimson clover + winter rye."},
    {"type":"bullet","text":"Establish forage / graze cover crops: Choose palatable species. Field peas, oats, and clovers are all good forage species. Suggested mix: Oats + field peas + crimson clover."},
    {"type":"h2","text":"Seeding Rates — Don't Overthink It"},
    {"type":"p","text":"Seeding rates matter, but they don't need to be precise to three decimal places. Here are simple guidelines:"},
    {"type":"bullet","text":"Winter rye: 60–90 lbs/acre drilled; 90–120 lbs broadcast."},
    {"type":"bullet","text":"Oats: 60–80 lbs/acre drilled; 80–100 lbs broadcast."},
    {"type":"bullet","text":"Hairy vetch: 15–25 lbs/acre."},
    {"type":"bullet","text":"Crimson clover: 15–20 lbs/acre."},
    {"type":"bullet","text":"Daikon radish: 5–8 lbs/acre."},
    {"type":"bullet","text":"Field peas: 40–60 lbs/acre."},
    {"type":"p","text":"When mixing species, reduce each individual rate by roughly 30–40%. When broadcasting (common when overseeding into standing crops), increase rates by 25–30%."},
    {"type":"h2","text":"Termination Timing and Methods"},
    {"type":"p","text":"Termination is where a lot of first-timers wait too long. The cover crop sets seed, becomes harder to kill, and creates a bigger mess than planned. Terminate before seed set. That's the main rule."},
    {"type":"bullet","text":"Mowing: Low-cost and practical for smaller fields. Works well for small-stemmed covers like clover and oats. Less effective on thick stems like sorghum-sudan or mature rye."},
    {"type":"bullet","text":"Grazing: The best termination method when it fits your operation. Livestock do the work, deposit nutrients, and you get a forage benefit. Works well for palatable mixes with field peas, oats, and clovers."},
    {"type":"bullet","text":"Rolling/crimping: No-till termination popular with organic producers. Requires the cover to be at or past anthesis (flowering) to work effectively. Creates a heavy mulch layer that suppresses weeds."},
    {"type":"bullet","text":"Herbicide termination: Fastest and most reliable. Glyphosate at label rate terminates most cover crops cleanly. Useful when weather or timing is tight."},
    {"type":"h2","text":"What to Avoid Your First Year"},
    {"type":"bullet","text":"Over-complicating the mix. Eight-species mixes sound impressive. Three to five species is plenty."},
    {"type":"bullet","text":"Planting too late. Cool-season covers need 4–6 weeks of growth before hard frost."},
    {"type":"bullet","text":"Forgetting about termination before you plant. Know how you're going to kill it before it goes in the ground."},
    {"type":"bullet","text":"Skipping inoculant for legumes. Legumes need specific rhizobia bacteria to fix nitrogen. If you haven't grown legumes in a field before, inoculate the seed. It's cheap and makes a real difference."},
    {"type":"h2","text":"A Sample Starter Mix for Pasture Improvement"},
    {"type":"p","text":"If you're improving a pasture and want a reliable first-year mix, here's a simple starting point:"},
    {"type":"bullet","text":"Winter rye: 40 lbs/acre"},
    {"type":"bullet","text":"Crimson clover: 10 lbs/acre"},
    {"type":"bullet","text":"Daikon radish: 4 lbs/acre"},
    {"type":"p","text":"Broadcast or drill in late summer to early fall. Inoculate the clover. Graze or mow before rye heads out in spring. This mix covers your bases: the rye adds biomass and winter cover, the clover fixes nitrogen and improves forage quality, and the radish breaks up any compaction."},
    {"type":"cta","text":"→ Nature's Seed carries cover crop seed and clover seed with real variety options — not just commodity blends. Browse cover crop seed mixes at naturesseed.com."},
  ]
},


# ── ARTICLE 5 ──────────────────────────────────────────────────────────────
{
  "title": "Frost Seeding Clover Into an Existing Pasture",
  "category": "Practical How-To",
  "read_time": "5 min read",
  "meta": "Frost seeding clover is the lowest-cost pasture improvement available — no tillage, no equipment, just seed and timing. Here's how to do it right.",
  "slug": "05-frost-seeding-clover",
  "body": [
    {"type":"p","text":"If you only do one thing to improve your pastures this year, make it frost seeding clover. It costs almost nothing, requires no equipment you don't already have, and works in the middle of winter when you have nothing else to do."},
    {"type":"p","text":"Frost seeding is the practice of broadcasting clover seed directly onto existing pasture in late winter — typically February through early March in most of the country — and letting freeze-thaw cycles work the seed into the soil naturally. No tillage. No drill. No seedbed prep."},
    {"type":"p","text":"The biology that makes it work: as temperatures swing above and below freezing, the soil surface heaves slightly. That movement creates micro-pockets that capture seed and pull it into contact with the soil. When spring arrives and temperatures climb, the clover germinates into an already-established pasture."},
    {"type":"h2","text":"Why Clover Specifically"},
    {"type":"p","text":"Clover is the ideal frost-seeding candidate for three reasons."},
    {"type":"p","text":"Small seed size. Clover seed is tiny — small enough to work into soil crevices created by freeze-thaw movement. Large-seeded species like grasses need actual soil disturbance to establish reliably. Clover doesn't."},
    {"type":"p","text":"Nitrogen fixation. Red and white clovers fix 80–150 lbs of atmospheric nitrogen per acre per year under good conditions. That's fertilizer your pasture produces itself. Every acre of established clover is an acre that needs less purchased nitrogen input."},
    {"type":"p","text":"Forage quality. Clover is highly digestible and palatable. Adding 20–30% clover to a grass-dominated stand measurably improves average daily gain in cattle and reduces the hay window in spring by improving early-season forage quality."},
    {"type":"h2","text":"Timing: The One Variable That Matters Most"},
    {"type":"p","text":"Frost seeding works because of freeze-thaw cycling. Seed too early (mid-winter, when the ground stays frozen) and seed just sits on frozen ground waiting for spring. Seed too late (when soil has thawed and grass is actively growing) and you lose the mechanical advantage of heaving soil."},
    {"type":"p","text":"The target window is when nighttime temperatures are still dropping below freezing but daytime temperatures are climbing above 32°F. In most of the Midwest and Northeast, that's mid-February through early March. In the upper Midwest or mountain states, it shifts two to four weeks later."},
    {"type":"p","text":"You want the last consistent freeze-thaw cycles of winter — not the dead of winter, not spring."},
    {"type":"h2","text":"Species Selection and Seeding Rates"},
    {"type":"p","text":"For most operations, the choice comes down to red clover versus white clover versus a mix of both."},
    {"type":"bullet","text":"Red clover: Faster-establishing, taller, higher biomass. Good for hay fields and pastures where you want quick-establishing forage. Short-lived (2–3 years) but reseeds if allowed to flower occasionally. Seeding rate: 6–10 lbs/acre."},
    {"type":"bullet","text":"White clover: Lower-growing, persistent, better for continuously grazed pastures. Spreads by stolons and can persist indefinitely in a well-managed stand. Seeding rate: 2–4 lbs/acre."},
    {"type":"bullet","text":"Mix: 4–6 lbs red + 2 lbs white covers both bases — quick establishment from red clover, long-term persistence from white."},
    {"type":"p","text":"Inoculate your seed if you haven't grown clover in the field in the past three to five years. Clover needs specific Rhizobium bacteria to fix nitrogen. The bacteria occur naturally in soils with clover history, but if that history is absent or uncertain, pre-inoculated seed or a powder inoculant applied at seeding is cheap insurance."},
    {"type":"h2","text":"Pasture Conditions That Set You Up for Success"},
    {"type":"p","text":"Frost seeding doesn't work equally well in every pasture. Success rate is highest when:"},
    {"type":"bullet","text":"The existing stand is thin or patchy. Thick, established grass will outcompete clover seedlings. Frost seeding works best where there are gaps in the canopy."},
    {"type":"bullet","text":"Soil pH is above 6.2. Clover struggles to establish and fix nitrogen in acidic soils. If your pH is below 6.0, lime the field first. Frost seeding into low-pH ground is money wasted."},
    {"type":"bullet","text":"Spring grazing is managed carefully. Keep livestock off the field until clover reaches 4–6 inches in height after germination. A single early grazing pass before establishment can wipe out your stand."},
    {"type":"h2","text":"What to Expect"},
    {"type":"p","text":"Frost-seeded clover doesn't look impressive until late spring. Germination is slow and the seedlings are tiny. Don't panic at slow establishment — this is normal. By July of the seeding year, a successful stand will be visible and providing real forage contribution."},
    {"type":"p","text":"Establishment success rates for frost seeding run 60–80% in good conditions. It's not perfect. But it's cheap enough that you can repeat it two years running if the first year is mediocre and still come out ahead compared to drilling."},
    {"type":"cta","text":"→ Nature's Seed carries red clover, white clover, and mixed clover seed with inoculant options. Farm-direct seed at honest rates — browse clover seed at naturesseed.com."},
  ]
},

# ── ARTICLE 6 ──────────────────────────────────────────────────────────────
{
  "title": "Renovating a Tired Hobby Farm Pasture in One Season",
  "category": "Practical How-To",
  "read_time": "9 min read",
  "meta": "Overseeding vs. full renovation for a depleted hobby farm pasture — how to diagnose which your ground needs, and what to plant to turn it around in one season.",
  "slug": "06-renovating-hobby-farm-pasture",
  "body": [
    {"type":"p","text":"Most small farms inherit a tired pasture. Maybe it was neglected, grazed too hard, or just never properly established to begin with. Thin stands, bare patches, weeds moving in, and productivity declining year by year."},
    {"type":"p","text":"The good news: a depleted pasture is fixable, usually within one growing season. The key is diagnosing how bad it actually is — because the answer determines whether you overseed into what's there or start fresh."},
    {"type":"h2","text":"The Diagnosis: Overseeding vs. Full Renovation"},
    {"type":"p","text":"Walk your pasture and do a rough estimate of your existing stand coverage. This determines your path."},
    {"type":"bullet","text":"More than 50% desirable grass/legume cover: Overseed. The existing stand is worth saving and competing against weeds. You can improve it significantly without killing it off."},
    {"type":"bullet","text":"25–50% desirable cover: Borderline. Overseeding can work but results are inconsistent. The weed pressure and competition from existing plants may limit establishment. Consider a light herbicide renovation before overseeding."},
    {"type":"bullet","text":"Less than 25% desirable cover: Full renovation. The existing stand isn't worth saving. Killing it and starting over gives you a cleaner seedbed and better establishment rates."},
    {"type":"p","text":"The honest calculus is cost vs. reliability. Overseeding is cheaper and less disruptive. Full renovation is more expensive upfront but gives you a predictable outcome. If you're below 30–35% desirable cover, the extra cost of renovation often pays for itself in first-year forage production."},
    {"type":"h2","text":"Overseeding a Tired Stand"},
    {"type":"p","text":"Overseeding works by introducing new seed into a thinned existing stand. For it to succeed, you need to get the seed into contact with bare soil — not on top of a mat of thatch or dense sod."},
    {"type":"h3","text":"Step 1 — Stress the existing stand first"},
    {"type":"p","text":"Graze the pasture hard in early spring or late fall, just before you plan to overseed. Get the existing grass short — 2 to 3 inches. This reduces competition for the new seedlings and exposes the soil surface. If grazing isn't an option, mow it low."},
    {"type":"h3","text":"Step 2 — Address the soil first"},
    {"type":"p","text":"Run a soil test before you spend money on seed. pH below 6.0 will cost you your clover. Compaction will cost you establishment depth. Lime if needed. This is the step most hobby farm operators skip, and it explains why their overseeding fails."},
    {"type":"h3","text":"Step 3 — Drill or use a slit seeder"},
    {"type":"p","text":"A no-till drill or a slit seeder (widely available for rent at farm co-ops and equipment dealers) places seed directly into the soil without full tillage. This is far superior to broadcasting — you get consistent seed-to-soil contact, better germination, and higher establishment rates. Budget $50–100/acre for rental if you don't own equipment. It's worth it."},
    {"type":"h3","text":"Step 4 — Manage grazing after seeding"},
    {"type":"p","text":"This is where most renovations fail. New seedlings need at least 60–90 days to develop root systems before they can withstand grazing pressure. Fence animals out completely if possible, or use temporary electric fence to protect newly seeded areas."},
    {"type":"h2","text":"Full Renovation: When to Burn It Down and Start Over"},
    {"type":"p","text":"Full renovation means killing the existing stand and establishing a new one from scratch. Done right, it gives you a clean seedbed and maximum flexibility to choose the right species mix for your goals."},
    {"type":"h3","text":"Killing the old stand"},
    {"type":"p","text":"The most reliable method is a full-rate glyphosate application when the existing grass is actively growing. Apply when temperatures are above 55°F and the stand is green and growing. A single application at label rate terminates most cool-season grass pastures cleanly. Allow 10–14 days after application before tillage or seeding."},
    {"type":"p","text":"For organic operations or those who prefer to avoid herbicide: repeated intensive tillage (two to three passes) in summer will desiccate most stands if timing aligns with hot, dry conditions. It takes longer and more passes than chemical termination but achieves a similar result."},
    {"type":"h3","text":"Seedbed preparation"},
    {"type":"p","text":"One to two passes with a disk or field cultivator creates a workable seedbed. You want firm, fine soil — not fluffy or cloddy. A culti-packer or roller after seeding improves seed-to-soil contact."},
    {"type":"h3","text":"Species selection for renovation"},
    {"type":"p","text":"The renovation decision is also an opportunity to be intentional about what goes back in the ground. For hobby farm pastures, a practical renovation mix includes:"},
    {"type":"bullet","text":"A perennial base grass: Orchardgrass for shade tolerance, tall fescue for durability, or timothy for hay quality. Choose based on your primary use and region."},
    {"type":"bullet","text":"A legume component: Red clover at 4–6 lbs/acre for quick nitrogen contribution, white clover at 2 lbs/acre for persistence, or birdsfoot trefoil for bloat-free grazing if that's a concern."},
    {"type":"bullet","text":"Optional: A small amount of chicory or plantain (2 lbs/acre) for mineral diversity and palatability."},
    {"type":"h2","text":"Timing the Renovation"},
    {"type":"p","text":"For cool-season grasses (the majority of pasture species in the US), two windows work best:"},
    {"type":"bullet","text":"Late summer seeding (August–September): Best results in most of the country. Cooler temperatures reduce weed competition. Rain reliability improves. Seedlings establish through fall and are ready to graze the following spring."},
    {"type":"bullet","text":"Spring seeding (March–May): Higher weed competition risk, but works when fall seeding isn't possible. Use a companion crop like oats to reduce weed pressure and provide early forage."},
    {"type":"p","text":"Late summer is the better window when you have a choice."},
    {"type":"h2","text":"Budget Reality for Hobby Farm Scale"},
    {"type":"p","text":"For a 10-acre renovation: seed ($150–250), herbicide ($80–120), equipment rental ($400–600 for drill and light tillage), and lime if needed ($150–200 for ag lime delivered). All-in budget for a 10-acre full renovation: $800–1,200, depending on region and what equipment you own."},
    {"type":"p","text":"That works out to $80–120 per acre for a pasture you'll use for the next decade. Against the cost of hay or the cost of a nutritionist explaining why your livestock aren't gaining, it's one of the better investments on a small farm."},
    {"type":"cta","text":"→ Nature's Seed carries perennial grass and legume seed suited for pasture renovation — orchard grass, tall fescue, clovers, and mixed pasture blends for hobby farms through larger operations. Browse pasture seed at naturesseed.com."},
  ]
},

# ── ARTICLE 7 ──────────────────────────────────────────────────────────────
{
  "title": "Stockpile Grazing: How to Save on Hay When You Have the Right Grass",
  "category": "Practical How-To",
  "read_time": "6 min read",
  "meta": "Stockpile grazing with tall fescue is one of the most cost-effective winter feeding strategies for small operations. Here's how the timing and management works.",
  "slug": "07-stockpile-grazing-save-on-hay",
  "body": [
    {"type":"p","text":"Hay is expensive. Cutting, baling, storing, and feeding hay represents one of the largest controllable costs in a cow-calf or small ruminant operation. And every bale your animals graze off the pasture instead of eating from a ring feeder is money that stays in your pocket."},
    {"type":"p","text":"Stockpile grazing is the practice of accumulating — stockpiling — grass growth in late summer and fall, then grazing it during winter months when it would otherwise be unavailable. Done well, it can extend your grazing season by 60–90 days and reduce hay consumption by 30–50% on a typical operation."},
    {"type":"p","text":"The grass that makes this possible is tall fescue. And the reason comes down to cold tolerance, protein retention, and management timing."},
    {"type":"h2","text":"Why Fescue Is the Stockpile Grass"},
    {"type":"p","text":"Most cool-season grasses decline in quality as they mature and dry down into winter. Tall fescue is different. It retains crude protein levels — often 12–16% — even after frost and light snow cover. Freeze-thaw cycling actually improves its palatability by reducing some of the compounds that make fresh fescue less appealing to livestock. By the time you're grazing it in December and January, the quality is often better than it was in September."},
    {"type":"p","text":"Orchardgrass and bluegrass can also be stockpiled, but results are more variable and quality declines faster. Bermudagrass works well in the South. But for the broadest geographic range and most consistent results, tall fescue is the stockpile standard."},
    {"type":"p","text":"One note on fescue variety: endophyte-infected tall fescue can cause toxicosis in cattle, particularly in late gestation and finishing animals. If you're establishing new fescue for stockpiling, use a novel-endophyte variety (like MaxQ) or endophyte-free seed. Existing fescue stands should be assessed — many older fields in the transition zone carry the toxic endophyte, and the negative effects on animal performance can offset your hay savings."},
    {"type":"h2","text":"The Stockpile Accumulation Window"},
    {"type":"p","text":"The process is straightforward, but timing is everything."},
    {"type":"p","text":"Close the field to grazing in late August to early September — ideally when you have 60 to 75 days before your expected first hard frost. This is your accumulation window. The fescue grows during the relatively warm days of late summer and early fall, accumulating significant dry matter."},
    {"type":"p","text":"Apply nitrogen in the first week of the accumulation period. Fifty pounds of actual nitrogen per acre is the standard recommendation. This single fertilizer application pays for itself many times over in additional stockpile growth. At current urea prices, you're spending $30–50/acre to potentially grow an additional 1,000–1,500 lbs of dry matter per acre."},
    {"type":"p","text":"Don't touch the field — no grazing, no clipping — from closure until you're ready to graze in late fall or early winter."},
    {"type":"h2","text":"How to Graze It Efficiently"},
    {"type":"p","text":"Strip grazing is the method that makes stockpile grazing work at its best. Rather than opening the entire stockpile field at once, use temporary electric fence to give animals access to a fresh strip every few days."},
    {"type":"p","text":"Why this matters: cattle are wasteful grazers when given large areas. They'll graze the most accessible sections repeatedly, trample and waste a significant portion of the stockpile, and you'll burn through your 90-day supply in 30 days. Strip grazing reduces waste by 40–60% and forces uniform utilization across the field."},
    {"type":"p","text":"Practical setup: run a single strand of polywire parallel to a fence line and move it every three to five days based on animal numbers and forage availability. A step-in post every 20 feet, a reel, and a single wire energizer is all the equipment you need. Cost: under $200 to outfit a field."},
    {"type":"h2","text":"How Much Stockpile Do You Need?"},
    {"type":"p","text":"A rough calculation: a mature cow needs approximately 2–2.5% of her body weight in dry forage per day. A 1,200-lb cow requires 24–30 lbs of dry matter daily. Well-managed stockpile can yield 2,000–3,500 lbs of dry matter per acre depending on fertility, rainfall, and variety."},
    {"type":"p","text":"For a 30-cow herd you want to carry 60 extra days without hay: 30 cows × 27 lbs/day × 60 days = 48,600 lbs dry matter needed. At 2,500 lbs/acre stockpile, that's roughly 20 acres. Adjust based on your specific numbers."},
    {"type":"h2","text":"What If You Don't Have Fescue?"},
    {"type":"p","text":"If your pastures are mostly bluegrass, orchardgrass, or native species, you have options. You can overseed a portion of your acreage with tall fescue this fall specifically for stockpile use. Fescue establishes readily in a prepared seedbed or with a no-till drill into existing pasture."},
    {"type":"p","text":"Alternatively, annual ryegrass can be used for late-season stockpile in milder climates (zones 6 and warmer). It won't match fescue's cold-hardiness, but in areas where temperatures stay above 15–20°F through December, annual ryegrass provides excellent late-season forage quality."},
    {"type":"cta","text":"→ Nature's Seed carries tall fescue varieties — including novel-endophyte options — and annual ryegrass for stockpile establishment. Farm-direct seed ready to ship. Browse at naturesseed.com."},
  ]
},

# ── ARTICLE 8 ──────────────────────────────────────────────────────────────
{
  "title": "Why Mycorrhizae Matter for Pasture Establishment (And How to Use Them)",
  "category": "Practical How-To",
  "read_time": "5 min read",
  "meta": "Mycorrhizal fungi dramatically improve pasture and cover crop establishment — here's the soil biology behind it and a practical guide to inoculating at seeding.",
  "slug": "08-mycorrhizae-pasture-establishment",
  "body": [
    {"type":"p","text":"Most pasture seeding guides tell you to focus on seedbed prep, seeding rate, and moisture. They're not wrong — those things matter. But they skip something that can double your establishment success rate and have a lasting impact on how your pasture performs for years afterward."},
    {"type":"p","text":"Mycorrhizal fungi. Specifically, arbuscular mycorrhizal fungi (AMF) — the most common type, and the one that colonizes grasses, legumes, and most pasture species."},
    {"type":"h2","text":"What Mycorrhizal Fungi Actually Do"},
    {"type":"p","text":"Mycorrhizal fungi form a symbiotic relationship with plant roots. The fungal filaments — called hyphae — extend far beyond the root zone, reaching pockets of phosphorus, water, and micronutrients that the plant's own roots can't access. In exchange, the plant supplies the fungi with sugars produced through photosynthesis."},
    {"type":"p","text":"The practical effect for a establishing pasture: plants colonized by mycorrhizal fungi have a dramatically larger effective root system. Under drought conditions, colonized seedlings can access moisture from a soil volume 10–100 times larger than their actual root zone. During the vulnerable establishment period — when seedlings are small and root systems shallow — that difference is often the margin between success and failure."},
    {"type":"p","text":"Beyond establishment, mycorrhizal fungi also improve phosphorus uptake in low-P soils, improve soil structure through glomalin production (a glue-like protein that binds soil aggregates), and make the plant more resilient to stress from drought, temperature swings, and pathogen pressure."},
    {"type":"h2","text":"The Problem: Many Farmed Soils Are Mycorrhizal Deserts"},
    {"type":"p","text":"Native soils with undisturbed plant communities are rich in mycorrhizal spore banks — dormant spores that germinate when roots pass nearby. Farmed soils are often depleted. Frequent tillage disrupts the fungal networks physically. Fumigation kills spore banks. High phosphorus fertility reduces plant incentive to maintain mycorrhizal associations (plants outsource nutrient acquisition to fungi only when they're nutrient-limited; provide all their phosphorus artificially and they don't bother)."},
    {"type":"p","text":"This is why establishing a new pasture in a recently tilled, fertilized field often produces disappointingly uneven results — particularly in the first two to three years. The biology isn't there to support robust establishment."},
    {"type":"h2","text":"How to Inoculate at Seeding"},
    {"type":"p","text":"The practical fix is straightforward: apply a mycorrhizal inoculant at seeding. This introduces a population of AMF spores directly to the seedling root zone, where they can colonize quickly without waiting for a dormant spore bank to reactivate."},
    {"type":"p","text":"Commercial mycorrhizal inoculants come in powder, granular, and liquid forms. For pasture seeding, the two most practical application methods are:"},
    {"type":"bullet","text":"Seed coating: Mix the dry inoculant powder with the seed before loading into the drill or spreader. Most products specify a ratio by weight. This gets the inoculant directly in contact with germinating seeds."},
    {"type":"bullet","text":"In-furrow granular: Apply granular inoculant through the fertilizer hopper on a grain drill, placing it directly in the seed trench. More precise placement than seed coating, and no concern about inoculant viability if there's a delay between coating and planting."},
    {"type":"h2","text":"Which Pasture Species Respond Best"},
    {"type":"p","text":"All common cool-season pasture grasses form mycorrhizal associations: orchardgrass, tall fescue, bluegrass, ryegrass. Legumes — clovers, alfalfa, trefoil — also form AMF associations, separate from and in addition to their rhizobial nitrogen-fixing symbiosis. Both are worth supporting."},
    {"type":"p","text":"One exception: brassicas (turnip, radish, canola, kale) do NOT form mycorrhizal associations. Don't waste inoculant on brassica-only stands."},
    {"type":"h2","text":"Managing for Mycorrhizal Health After Establishment"},
    {"type":"p","text":"Inoculating at seeding is just the start. How you manage the pasture afterward determines whether the fungal community builds or collapses."},
    {"type":"bullet","text":"Avoid unnecessary tillage. Every tillage pass severs hyphal networks. In an established pasture, this isn't usually an issue — but renovation work disrupts the network and is another reason to overseed rather than till when possible."},
    {"type":"bullet","text":"Keep living roots in the ground. Fungi depend on plant hosts for energy. Bare soil periods — summer burn-down before renovation, extreme overgrazing — allow the fungal population to decline."},
    {"type":"bullet","text":"Use phosphorus inputs carefully. High soil phosphorus suppresses mycorrhizal associations. Apply phosphorus based on soil test results, not standard rates. Over-application is expensive and biologically counterproductive."},
    {"type":"cta","text":"→ Nature's Seed seed experts can help you choose the right pasture mix and talk through establishment practices for your region. Browse pasture and cover crop seed — or ask a seed specialist — at naturesseed.com."},
  ]
},

# ── ARTICLE 9 ──────────────────────────────────────────────────────────────
{
  "title": "Silvopasture: Growing Forage Under Trees Without Killing Either",
  "category": "Niche Deep Dives",
  "read_time": "10 min read",
  "meta": "Silvopasture integrates trees and livestock pasture for long-term productivity. Species selection, spacing, and shade management for small and medium operations.",
  "slug": "09-silvopasture-forage-under-trees",
  "body": [
    {"type":"p","text":"Silvopasture is one of the older land use systems in agriculture, predating modern row cropping by centuries. Farmers in Europe and Latin America have grazed livestock in wooded settings for generations. In North America, it largely disappeared as agriculture industrialized. It's coming back now — partly because of carbon payment programs, partly because a growing number of ranchers have rediscovered that trees and livestock can coexist profitably."},
    {"type":"p","text":"The pitch for silvopasture on a small or medium operation isn't primarily about carbon credits. It's about stacking productive uses on the same acre: timber or nut crop income layered on top of grazing income, with a microclimate under the canopy that can actually improve forage quality compared to open pasture in hot summer months."},
    {"type":"p","text":"The challenge: getting the tree and forage components to coexist without the trees shading out the grass or the livestock girdling the trees. That's a management and design problem, and it's solvable."},
    {"type":"h2","text":"What Silvopasture Actually Looks Like"},
    {"type":"p","text":"There are three basic configurations, and the right one depends on what you're starting with."},
    {"type":"h3","text":"Trees planted into existing pasture"},
    {"type":"p","text":"The most common approach for small operations starting from scratch. Establish rows of trees across existing pasture at wide spacing — 30 to 60 feet between rows depending on species — and manage the alleys between rows as normal pasture. Tree density is low enough that ground-level light is maintained for forage production."},
    {"type":"h3","text":"Pasture established under existing trees"},
    {"type":"p","text":"Common on farms that have existing woodland or established shelter belts. Clear the understory brush, open up canopy where needed for adequate light, and establish shade-tolerant forage species in the openings. This approach is cheaper than planting new trees and produces faster results."},
    {"type":"h3","text":"Thinned forest conversion"},
    {"type":"p","text":"Take an existing woodlot, thin it to a target basal area that allows enough light for forage growth (typically 40–60% canopy cover), and establish forage in the openings. This is the most complex approach but also the one with the most potential for existing forested land."},
    {"type":"h2","text":"Light: The Variable Everything Else Depends On"},
    {"type":"p","text":"Forage grasses need light. Most cool-season grasses require at least 50–60% of full sunlight to maintain productive stands. Below that threshold, grass thins out, weed species tolerant of shade move in, and forage quality declines."},
    {"type":"p","text":"The design goal of silvopasture is to maintain canopy cover in the 30–50% range — enough shade to provide livestock comfort on hot summer days and buffer forage quality, but not so much that forage production collapses. Getting there requires attention to tree row orientation (east-west rows create more uniform shade distribution), spacing, and long-term pruning management."},
    {"type":"p","text":"As trees mature and canopy closes, you'll need to either manage row width through pruning or accept that forage production gradually shifts from light-demanding grasses to shade-tolerant species and eventually to a forest understory system."},
    {"type":"h2","text":"Tree Species Selection"},
    {"type":"p","text":"The right tree species depends on your goals, region, and market access. General categories:"},
    {"type":"h3","text":"Timber species"},
    {"type":"p","text":"Black walnut, white oak, and black locust are common choices in the eastern US. Black locust fixes nitrogen, grows fast, and produces rot-resistant wood — it's arguably the best multifunctional silvopasture tree in humid climates east of the Mississippi. Black walnut allelopathic compounds can suppress some pasture species; keep grass rows 10–12 feet from walnut trunks and focus on allelopathy-tolerant species (tall fescue handles it reasonably well)."},
    {"type":"h3","text":"Nut and fruit species"},
    {"type":"p","text":"Chestnut and pecan for nut production. Chestnuts in particular work well in silvopasture — fast-growing, non-allelopathic, and the nuts make excellent livestock feed for the portion that falls and isn't harvested. Apples and pears in orchard configurations with livestock grazing between rows (a practice common in old orchards) is a simpler, lower-establishment-cost version of the same idea."},
    {"type":"h3","text":"Fodder trees"},
    {"type":"p","text":"Mulberry and willow cut-and-carry systems have gotten renewed attention as a direct livestock feed supplement. Mulberry leaves have crude protein levels of 15–20%, comparable to alfalfa, and can be harvested by coppicing — cutting back to a stump to produce fresh growth repeatedly."},
    {"type":"h2","text":"Forage Species for Shade"},
    {"type":"p","text":"Not all pasture species tolerate partial shade equally well. Performance ranking under moderate shade (30–50% canopy cover):"},
    {"type":"bullet","text":"Good tolerance: Orchardgrass, endophyte-free tall fescue, white clover, birdsfoot trefoil."},
    {"type":"bullet","text":"Moderate tolerance: Timothy, red clover, annual ryegrass."},
    {"type":"bullet","text":"Poor tolerance: Bermudagrass, most warm-season grasses, alfalfa."},
    {"type":"p","text":"Orchardgrass is the standout performer in silvopasture systems across the temperate US. It was named for its original habitat — orchards — and its shade tolerance reflects that history. It's productive, palatable, and persistent under the kind of partial shade a well-managed silvopasture system provides."},
    {"type":"h2","text":"Protecting Young Trees From Livestock"},
    {"type":"p","text":"This is where most silvopasture attempts fail in years one through five. Young trees are defenseless against livestock. Cattle will rub, browse, and girdle a sapling in a single grazing event."},
    {"type":"p","text":"Tree protection options in roughly increasing cost:"},
    {"type":"bullet","text":"Individual tree tubes (treeshelters): Plastic tubes that protect the trunk and accelerate early growth. Cost: $3–6 per tree. Works well for smaller diameter planting stock. Provides 3–5 years of protection before trees outgrow them."},
    {"type":"bullet","text":"Wire cages: 4–5 feet of hardware cloth or cattle panel formed into a cylinder. More durable than treeshelters and better for browsing pressure from sheep and goats. Cost: $8–15 per tree."},
    {"type":"bullet","text":"Temporary electric fence exclusion: Fence off entire tree rows until trees reach a size that can withstand livestock contact (typically 3-inch trunk diameter or larger). Remove fencing once trees are established. Higher upfront cost but flexible and reusable."},
    {"type":"h2","text":"Carbon and Cost-Share Opportunities"},
    {"type":"p","text":"Silvopasture is one of the practices explicitly supported under USDA EQIP (Environmental Quality Incentives Program) cost-sharing. Payments vary by state but can offset 50–75% of tree establishment costs. Your local NRCS office is the starting point for applications."},
    {"type":"p","text":"Carbon markets are increasingly offering payments for silvopasture establishment through programs like the USDA's voluntary carbon market initiatives and private aggregators. Per-acre payments range widely ($10–50/acre/year depending on the program and verification standards). The market is immature and monitoring requirements can be burdensome for small operations, but it's worth understanding what's available."},
    {"type":"cta","text":"→ Orchardgrass and white clover are the core of most silvopasture forage mixes. Nature's Seed carries both, along with regionally appropriate shade-tolerant species. Browse pasture seed at naturesseed.com."},
  ]
},

# ── ARTICLE 10 ──────────────────────────────────────────────────────────────
{
  "title": "Pollinator Forage for Beekeepers With Acreage: Beyond Clover",
  "category": "Niche Deep Dives",
  "read_time": "7 min read",
  "meta": "Building a season-long pollinator forage calendar for beekeepers with acreage — species selection, bloom timing, and how to sequence plantings for continuous nectar flow.",
  "slug": "10-pollinator-forage-beekeepers",
  "body": [
    {"type":"p","text":"Most beekeeping guides tell you to plant clover and call it done. That's not wrong — clovers are excellent pollinator plants and easy to establish on acreage. But if you have more than a few acres to work with, you can build something significantly better: a managed forage calendar that provides nectar and pollen from spring through hard frost, extending your colony's productive season and reducing supplemental feeding."},
    {"type":"p","text":"The goal isn't more species for the sake of diversity. It's strategic bloom sequencing — different plants flowering at different windows so your colonies always have something working."},
    {"type":"h2","text":"The Calendar Problem"},
    {"type":"p","text":"Most forage landscapes have boom-and-bust patterns. There's an intense spring flush — dandelions, fruit tree bloom, early clover — followed by a mid-summer dearth as temperatures spike and most plantings finish blooming. Then a fall flow from asters and goldenrod, if you're lucky."},
    {"type":"p","text":"The summer dearth is when colonies starve, swarm from stress, or simply fail to build winter reserves. If you have acreage, you can fill that gap with intentional plantings."},
    {"type":"h2","text":"Early Season (April–May)"},
    {"type":"p","text":"The goal in early season is building colony population ahead of the main nectar flow. Pollen matters as much as nectar here — bees need protein to raise brood."},
    {"type":"bullet","text":"Phacelia (Phacelia tanacetifolia): One of the best early-season bee plants in North America. Deep blue flowers, long bloom window of 4–6 weeks, extremely high nectar and pollen production. Reseeds aggressively if you let it go to seed. Annual. Seeding rate: 2–4 lbs/acre."},
    {"type":"bullet","text":"Crimson clover: Earlier-blooming than red or white clover, typically April–May in zones 6+. Excellent pollen source. Annual in most climates. Seeding rate: 10–15 lbs/acre."},
    {"type":"bullet","text":"Borage: Fast-establishing annual with a very long bloom window. Self-seeds readily. Good gap-filler in small plots."},
    {"type":"h2","text":"Main Flow Season (June–July)"},
    {"type":"p","text":"This is when most colonies make their surplus honey. The foundation:"},
    {"type":"bullet","text":"White clover (Trifolium repens): The gold standard for honey production. High nectar sugar concentration, exceptionally accessible flower depth for honeybees. Persistent perennial that spreads by stolons. Tolerates repeated mowing and grazing. The non-negotiable inclusion in any beekeeper's forage plan. Seeding rate: 2–4 lbs/acre."},
    {"type":"bullet","text":"Red clover: Higher biomass than white, blooms slightly later. Flower tube depth can limit nectar access for short-tongued honeybees — mixed results. Better for bumblebees and long-tongued species. Worth including but not the primary honey producer."},
    {"type":"bullet","text":"Sweet clover (Melilotus officinalis / M. alba): One of the best honey plants in North America when allowed to bloom freely. Biennial. Produces enormous biomass in year two before flowering and dying. Honey from sweet clover operations commands a premium. Note: sweet clover contains coumarin and can be problematic in hay — keep it in designated pollinator plots rather than multi-use hay fields."},
    {"type":"bullet","text":"Buckwheat: A warm-season annual that fills the midsummer gap where cool-season clovers have finished. Blooms intensely for 4–6 weeks in midsummer. Honey is dark and strongly flavored — valued by some markets, not others. Seeding rate: 40–50 lbs/acre. Plant after last frost."},
    {"type":"h2","text":"Late Season and Fall (August–October)"},
    {"type":"p","text":"Fall forage is critical for winter prep. Colonies need to build winter stores through September. Native plants are your best allies here."},
    {"type":"bullet","text":"Goldenrod (Solidago spp.): The most important fall nectar plant in North America east of the Rockies. Blooms August through October depending on species. Native perennial, spreads readily, requires zero maintenance once established. Often viewed as a weed; beekeepers know better."},
    {"type":"bullet","text":"Native asters (Symphyotrichum spp.): Bloom alongside and after goldenrod, extending the fall flow into October. Excellent both as nectar and pollen sources. Low-growing species like heath aster work well in wildflower mixes."},
    {"type":"bullet","text":"Phacelia as a fall planting: A second seeding of phacelia in August will bloom in September–October in warmer zones. Where timing works, it's one of the best late-season additions."},
    {"type":"h2","text":"Designing the Acreage Layout"},
    {"type":"p","text":"You don't need to dedicate every acre to pollinator forage. A practical approach for a 10–50 acre operation:"},
    {"type":"bullet","text":"Permanent perennial base: White clover overseeded into your general pasture provides the main-flow foundation. This doubles as livestock forage and doesn't require dedicated acres."},
    {"type":"bullet","text":"Dedicated annual strips: Rotated 1–3 acre plots for high-value annuals like phacelia, buckwheat, and sweet clover. Move these strips around your acreage on a 2–3 year rotation to maintain soil health and prevent stand decline."},
    {"type":"bullet","text":"Permanent native perennial patches: Establish goldenrod and native asters in edges, fence lines, and rough areas that aren't productive pasture. These areas contribute fall forage without competing with other land uses."},
    {"type":"h2","text":"What to Avoid"},
    {"type":"bullet","text":"Alfalfa: High nectar producer but flower mechanism restricts honeybee access. Better suited to leafcutter bees and bumbles."},
    {"type":"bullet","text":"Most ornamental flowers: Selected for appearance, not nectar production. Often double-petaled varieties that physically block bee access."},
    {"type":"bullet","text":"Treated seed: Neonicotinoid-coated seed in pollinator plots undermines the entire effort. Verify seed is untreated if your supplier doesn't specify."},
    {"type":"cta","text":"→ Nature's Seed carries white clover, sweet clover, crimson clover, and wildflower mixes including pollinator-specific blends. Farm-direct, no fillers. Browse at naturesseed.com."},
  ]
},

# ── ARTICLE 11 ──────────────────────────────────────────────────────────────
{
  "title": "Carbon Capture on a Small Farm: What's Real, What's Hype",
  "category": "Niche Deep Dives",
  "read_time": "8 min read",
  "meta": "Soil carbon programs for farms under 500 acres explained — what the science actually says, what payments are realistic, and whether it's worth the paperwork.",
  "slug": "11-carbon-capture-small-farm",
  "body": [
    {"type":"p","text":"Every few months a new carbon farming program gets announced with promises of meaningful income for small landowners. Most of them deliver less than the headline suggests. A few are worth the time to understand."},
    {"type":"p","text":"If you manage fewer than 500 acres and are wondering whether soil carbon sequestration is a real opportunity or a distraction, this is the honest breakdown."},
    {"type":"h2","text":"What the Science Actually Says"},
    {"type":"p","text":"Soils can sequester carbon. This is not disputed. When organic matter builds in soil — through cover crops, reduced tillage, improved grazing management — a portion of that carbon becomes stable in the soil for years to decades. This is measurable."},
    {"type":"p","text":"What's more complicated is the rate, permanence, and additionality of sequestration at farm scale. Here's where the science gets honest:"},
    {"type":"bullet","text":"Rates vary enormously. Well-managed pasture under rotational grazing can sequester 0.5–1.5 metric tons of CO2 equivalent per acre per year. Converting tilled cropland to no-till can sequester 0.1–0.5 tons/acre/year. These are ranges from real research — not maximums you'll achieve every year."},
    {"type":"bullet","text":"Carbon already in the soil doesn't stay there forever. A drought year, a policy change that reverts you to tillage, or a change in ownership can release sequestered carbon. This is the permanence problem that markets are still working out."},
    {"type":"bullet","text":"Additionality is a bureaucratic term that matters practically: markets pay for carbon sequestered beyond what would have happened anyway. If you were already no-tilling, you may not qualify for credit for that practice — even though your soil has more carbon than a tilled field."},
    {"type":"h2","text":"The Market Reality for Operations Under 500 Acres"},
    {"type":"p","text":"Carbon markets have three cost structures that create a problem for small operations:"},
    {"type":"p","text":"Monitoring, reporting, and verification (MRV) costs. Documenting soil carbon requires baseline soil sampling, periodic re-sampling, and third-party verification. Soil sampling alone costs $20–50 per sample point, and a rigorous baseline requires multiple points per field per soil type. These fixed costs don't scale down with acreage — they eat a disproportionate share of small-operation payments."},
    {"type":"p","text":"Market carbon prices. Voluntary carbon markets have been paying $10–25 per metric ton of CO2 equivalent for agricultural soil carbon. At 0.5 tons/acre/year sequestration (a reasonable middle estimate), that's $5–12.50 per acre per year before MRV costs."},
    {"type":"p","text":"The math for 100 acres: gross revenue of $500–1,250/year, minus MRV costs that can easily run $2,000–5,000/year for a rigorous program. Many small operations lose money in year one and two and don't break even until later in a long-term contract."},
    {"type":"p","text":"The economic threshold for standalone carbon revenue is roughly 300–500 acres under current market conditions. Below that, direct payment programs are better suited to small operations."},
    {"type":"h2","text":"What Actually Works for Small Operations"},
    {"type":"h3","text":"USDA RCPP and EQIP Programs"},
    {"type":"p","text":"The USDA's Regional Conservation Partnership Program (RCPP) and Environmental Quality Incentives Program (EQIP) provide cost-share and direct payments for soil health practices — cover cropping, no-till, rotational grazing fencing and water infrastructure. These aren't framed as carbon payments, but the practices they support are exactly the ones that sequester carbon."},
    {"type":"p","text":"The advantage over carbon markets: no MRV costs, no contract lock-in beyond the agreement period, and payments are based on practice implementation rather than measured carbon outcomes. A 50-acre farm can receive meaningful cost-share for cover crop seed and fencing through EQIP that a carbon market contract would never make economical."},
    {"type":"h3","text":"Aggregated Carbon Programs"},
    {"type":"p","text":"Some programs aggregate small producers under a single verification umbrella, spreading MRV costs across many farms. Indigo Agriculture, Ecosystem Services Market Consortium (ESMC), and several state-level programs have used this model. The economics improve compared to individual contracts, but additionality requirements still limit eligibility for farms already practicing no-till or cover cropping."},
    {"type":"h2","text":"The Honest Framing"},
    {"type":"p","text":"For most small farms, soil carbon sequestration is a co-benefit of good land management — not a primary revenue source. The practices that build carbon also reduce input costs, improve drought resilience, and maintain long-term productivity. Those operational benefits are real and accessible at any scale."},
    {"type":"p","text":"Carbon payments may become more meaningful as market standards mature and MRV costs come down through remote sensing and AI-assisted monitoring. The trajectory is positive. But in 2025, the realistic carbon revenue for a 100-acre farm is a rounding error compared to the operational benefits of the underlying practices."},
    {"type":"p","text":"Do the practices. Pursue the USDA cost-share. Evaluate carbon program eligibility honestly. Don't make land management decisions based on speculative carbon payment projections."},
    {"type":"h2","text":"Where to Learn More"},
    {"type":"bullet","text":"NRCS EQIP: Contact your local NRCS office. Applications open in fall for the following year in most states."},
    {"type":"bullet","text":"Soil Carbon Initiative (sciencebasedtargets.org): Overview of corporate demand for soil carbon that drives private market programs."},
    {"type":"bullet","text":"Rodale Institute: Long-term research on organic and regenerative system carbon outcomes."},
    {"type":"cta","text":"→ Cover crops and managed grazing are the highest-ROI soil carbon practices — and they don't require a carbon contract to be worthwhile. Nature's Seed carries cover crop mixes and pasture seed designed for real operations. Browse at naturesseed.com."},
  ]
},

# ── ARTICLE 12 ──────────────────────────────────────────────────────────────
{
  "title": "Multi-Species Cover Crop Mixes Explained: Why Single Species Are Leaving Performance on the Table",
  "category": "Niche Deep Dives",
  "read_time": "7 min read",
  "meta": "Multi-species cover crop mixes outperform single-species plantings on nearly every metric. Here's the science and practical design guide for building effective mixes.",
  "slug": "12-multi-species-cover-crop-mixes",
  "body": [
    {"type":"p","text":"Straight winter rye is still the most commonly planted cover crop in the US. It's cheap, it's reliable, and it works. But if you've been planting single-species covers for a few years and wondering whether there's more to get out of them, the answer is yes — and the path there is mixing species."},
    {"type":"p","text":"Multi-species cover crop mixes consistently outperform single-species plantings across biomass production, weed suppression, nitrogen contribution, soil biology, and cash crop yield response. This isn't theoretical — it's been studied extensively on working farms across multiple soil types and climates."},
    {"type":"h2","text":"The Functional Diversity Principle"},
    {"type":"p","text":"Different plant species do different things in the soil. Root architectures vary — some go deep, some stay shallow. Some produce fine roots that feed bacteria, others produce coarser roots that support fungal networks. Some plants exude sugars that feed specific microbial communities. Some fix nitrogen. Some scavenge nutrients from deeper horizons and bring them to the surface."},
    {"type":"p","text":"A single species optimizes for one function. A well-designed mix optimizes for several simultaneously."},
    {"type":"p","text":"The principle is similar to a work crew: one person who's excellent at one task is good, but a small team with complementary skills accomplishes more. A five-species cover crop mix with a grass, two legumes, a brassica, and a broadleaf covers biological nitrogen fixation, biomass production, compaction-breaking, nutrient scavenging, and surface soil biology in a single planting."},
    {"type":"h2","text":"What the Research Shows"},
    {"type":"p","text":"Studies from multiple university extension programs comparing single-species covers to diverse mixes have consistently found:"},
    {"type":"bullet","text":"Biomass: 3–5 species mixes produce 20–40% more total above-ground biomass than single-species plantings at comparable seeding costs."},
    {"type":"bullet","text":"Nitrogen contribution: Mixed legume-grass plantings deliver more plant-available nitrogen to the following cash crop than monoculture legume stands, because the grass component protects legume biomass from premature breakdown."},
    {"type":"bullet","text":"Weed suppression: Multi-species mixes create more complete canopy coverage, reducing weed emergence by 30–60% compared to single-species covers."},
    {"type":"bullet","text":"Soil biology: Diverse plantings support measurably higher fungal-to-bacterial ratios and greater microbial diversity in the soil."},
    {"type":"h2","text":"The Four Functional Groups to Design Around"},
    {"type":"h3","text":"Grasses (biomass and organic matter)"},
    {"type":"p","text":"High carbon-to-nitrogen ratio. Slow decomposition. Provides long-lasting organic matter addition. Also the primary weed-suppression driver through sheer biomass. Core options: winter rye, oats, sorghum-sudan (warm season), winter wheat."},
    {"type":"h3","text":"Legumes (nitrogen fixation)"},
    {"type":"p","text":"Low carbon-to-nitrogen ratio. Decompose quickly and release nitrogen into soil. Target 30–50% of your mix by species, not by weight — legumes are small-seeded. Core options: hairy vetch, crimson clover, field peas, sunn hemp (warm season)."},
    {"type":"h3","text":"Brassicas (compaction and nutrient scavenging)"},
    {"type":"p","text":"Taproots physically penetrate compaction layers. Scavenge nitrate from depth. Biofumigation effect from glucosinolates (natural fungicide activity) can suppress some root pathogens. Core options: daikon radish, turnip, rapeseed, kale."},
    {"type":"h3","text":"Broadleaf non-legumes (biology and diversity)"},
    {"type":"p","text":"Often the most overlooked group. Plants like sunflower, buckwheat, phacelia, and flax bring root exudate diversity that feeds different soil microbial communities. Some provide pollinator benefit. Small amounts (1–3 lbs/acre) add biological diversity without driving cost."},
    {"type":"h2","text":"Designing a Mix: Practical Framework"},
    {"type":"p","text":"Start with your primary goal and build around it."},
    {"type":"bullet","text":"Primary goal: nitrogen for next crop — Lead with legumes (40–50% of seeding rate). Add a supporting grass to protect legume biomass. Example: hairy vetch 15 lbs + winter rye 40 lbs + crimson clover 8 lbs."},
    {"type":"bullet","text":"Primary goal: biomass / organic matter — Lead with grasses. Example: winter rye 50 lbs + oats 30 lbs + hairy vetch 12 lbs + daikon radish 4 lbs."},
    {"type":"bullet","text":"Primary goal: forage / grazing — Include palatable species. Example: oats 50 lbs + field peas 30 lbs + crimson clover 10 lbs + turnip 3 lbs."},
    {"type":"bullet","text":"Primary goal: soil biology recovery — Maximize diversity across all four groups. Small amounts of many species. Example: winter rye 30 lbs + hairy vetch 12 lbs + crimson clover 8 lbs + daikon radish 4 lbs + phacelia 2 lbs + sunflower 3 lbs."},
    {"type":"h2","text":"Seeding Rate Math for Mixes"},
    {"type":"p","text":"The key principle: when mixing species, reduce individual species rates by 30–50% from their monoculture rate. The competition between species in a dense mix means each doesn't need its full rate to achieve adequate stand."},
    {"type":"p","text":"A five-species mix targeting 80–90 lbs/acre total isn't unusual. The cost often comes in similar to a high-rate monoculture because you're reducing each individual rate, even though you're buying more SKUs."},
    {"type":"h2","text":"Common Mistakes in Mix Design"},
    {"type":"bullet","text":"Too many species without functional reason. Eight species can outperform five, but not reliably. Adding species for complexity rather than function dilutes focus."},
    {"type":"bullet","text":"Ignoring termination timing differences. Species in a mix don't all terminate at the same time or with the same method. Winter rye at anthesis and immature crimson clover have different response rates to rolling. Design for how you're going to kill it."},
    {"type":"bullet","text":"Forgetting inoculant on legumes. Every legume species has specific rhizobia requirements. A hairy vetch + crimson clover + field pea mix needs inoculant that covers all three, or separate inoculants for each."},
    {"type":"bullet","text":"Mixing incompatible species. Brassicas and legumes can compete heavily in wet years. Dominant grasses can shade out small-seeded legumes if rates aren't balanced. Test new mixes on a small acreage before scaling."},
    {"type":"cta","text":"→ Nature's Seed carries the individual species you need to build custom cover crop mixes — grasses, legumes, brassicas, and specialty broadleafs. Farm-direct, no fillers, with seed experts who can help you dial in a mix for your ground. Browse cover crop seed at naturesseed.com."},
  ]
},
]


def slugify(slug):
    return slug.replace("/", "-").replace(" ", "-").lower()


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    for article in ARTICLES:
        doc = build_doc(article)
        filename = f"{article['slug']}.docx"
        path = os.path.join(OUTPUT_DIR, filename)
        doc.save(path)
        print(f"  Saved: {filename}")
    print(f"\nDone. {len(ARTICLES)} articles written to {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
